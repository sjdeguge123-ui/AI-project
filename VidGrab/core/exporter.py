# 多格式导出模块
# 支持 Markdown / HTML / Word / 图片(PNG) 四种格式，文件名用视频标题
"""把 Summary 渲染并导出为文件。

入口：
  export(summary, output, transcript, formats) -> List[Path]
  export_markdown / export_html / export_docx / export_image  单格式快捷入口

文件命名规则：{标题摘要}[-P{n}].{扩展名}
  - 标题中的非法文件名字符（\\/:*?"<>|）会被替换为下划线
  - 标题过长（>30 字）自动截断为「前 20 字 + …」的标题摘要形式
  - 合集视频追加 -P{n}（n 为选中集，1-based）；单P 视频不带 -P
  - 文件名始终带「摘要」二字，表明这是摘要文档
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

from datetime import date

from . import Summary, Transcript
from .config import OutputConfig
from .templates import render_summary_md


_PROJECT_ROOT = Path(__file__).resolve().parent.parent

# 文件名非法字符（Windows/Linux 通用）
_ILLEGAL_CHARS = re.compile(r'[\\/:*?"<>|]')


# 文件名长度阈值（用户 2026-07-20 指定）：标题部分超过此字数需要截断，
# 但「摘要」「-Pxx」等后缀仍保留。标题截断为「前 29 字 + 省略号」，刚好 30 字。
_TITLE_MAX = 30       # 标题部分超过 30 字 → 截断
_TITLE_DIGEST = 29    # 截断后保留前 29 字 + "…"，标题部分最多 30 字


def _safe_title(transcript: Transcript, mode_label: str = "") -> str:
    """生成安全的输出文件名（不含扩展名）。

    规则（用户 2026-07-20 指定）：
      - 标题部分不长（<= _TITLE_MAX）：用完整标题；
      - 标题部分过长（> _TITLE_MAX）：改为「前 _TITLE_DIGEST 字 + …」的标题摘要，
        标题部分最多 30 字；
      - 始终带「摘要」二字，表明这是摘要文档；
      - 合集：文件名追加 -P{n}（n 为选中集，1-based）；单P 视频不带 -P。
      - 追加模式标识（精简/详细/自定义），避免不同模式输出互相覆盖。
      - 注意："摘要"、"-Pxx"、"-精简" 等后缀不计入 30 字限制。

    示例：
      短标题单P-精简     -> 如何三天搞定深度学习摘要-精简.md
      长标题单P-详细     -> 从零开始学Blender建模渲染全流程详解…摘要-详细.md
      长标题合集P92-精简 -> 从零开始学Blender建模渲染全流程详解…摘要-精简-P92.md
    """
    title = transcript.title or f"{transcript.platform.value}_{transcript.video_id}"
    title = _ILLEGAL_CHARS.sub("_", title).strip().strip(".")

    # 标题过长则改为「标题摘要」（截断 + 省略号）
    if len(title) > _TITLE_MAX:
        title = title[:_TITLE_DIGEST].rstrip() + "…"

    # 始终带「摘要」标识；追加模式标识；合集（多P）追加 -P{n}，单P 视频不带 -P
    name = f"{title}摘要"
    if mode_label:
        name = f"{name}-{mode_label}"
    if transcript.is_collection:
        n = (transcript.page_index or 0) + 1
        name = f"{name}-P{n}"
    return name


def _resolve_save_dir(output: OutputConfig) -> Path:
    """决定落盘目录：默认项目 output/YYYYMMDD/；配置 save_path 时则 save_path/YYYYMMDD/。

    按日期分组便于用户后续整理、回看历史输出，与常见的下载/日志习惯一致。
    相对路径相对项目根；未配置则用项目下 output/YYYYMMDD/。
    """

    today = date.today().strftime("%Y%m%d")
    if output.save_path:
        d = Path(output.save_path)
        if not d.is_absolute():
            d = _PROJECT_ROOT / d
    else:
        d = _PROJECT_ROOT / "output"
    return d / today


def export_markdown(summary: Summary, output: OutputConfig, transcript: Transcript, mode_label: str = "") -> Path:
    """渲染成 Markdown 并写入 .md 文件，返回路径。"""

    text = render_summary_md(summary)
    save_dir = _resolve_save_dir(output)
    save_dir.mkdir(parents=True, exist_ok=True)
    path = save_dir / f"{_safe_title(transcript, mode_label)}.md"
    path.write_text(text, encoding="utf-8")
    return path


def _build_html(summary: Summary) -> str:
    """把 Summary 渲染成完整 HTML 字符串（供 HTML 导出和 PDF 导出复用）。"""

    import markdown as md

    md_text = render_summary_md(summary)
    html_body = md.markdown(md_text, extensions=["tables", "fenced_code"])
    # 基本样式：中文友好、表格美观、金句高亮
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{summary.title} - 摘要</title>
<style>
  body {{ font-family: -apple-system, "Microsoft YaHei", "PingFang SC", sans-serif; max-width: 800px; margin: 2em auto; padding: 0 1em; line-height: 1.8; color: #333; }}
  h2 {{ border-bottom: 2px solid #4a90d9; padding-bottom: 0.3em; color: #2c3e50; }}
  h3 {{ color: #34495e; margin-top: 1.5em; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
  th {{ background: #f5f7fa; font-weight: bold; }}
  tr:nth-child(even) {{ background: #fafbfc; }}
  strong {{ color: #e74c3c; }}
  blockquote {{ border-left: 4px solid #4a90d9; margin: 1em 0; padding: 0.5em 1em; background: #f8f9fa; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""


def export_html(summary: Summary, output: OutputConfig, transcript: Transcript, mode_label: str = "") -> Path:
    """渲染成 HTML 并写入 .html 文件，带基本样式，返回路径。"""

    html = _build_html(summary)
    save_dir = _resolve_save_dir(output)
    save_dir.mkdir(parents=True, exist_ok=True)
    path = save_dir / f"{_safe_title(transcript, mode_label)}.html"
    path.write_text(html, encoding="utf-8")
    return path


def export_pdf(summary: Summary, output: OutputConfig, transcript: Transcript, mode_label: str = "") -> Path:
    """渲染成 PDF 文件，返回路径。

    使用 reportlab 原生 Table + Paragraph 绘制，完全控制表格列宽与中文 CJK 折行，
    避开 xhtml2pdf 对长表格/长中文内容溢出截断的问题。字体用 reportlab 自带
    Adobe 标准简体中文字体 STSong-Light，零系统依赖，文本可复制/提取。
    """

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether

    save_dir = _resolve_save_dir(output)
    save_dir.mkdir(parents=True, exist_ok=True)
    path = save_dir / f"{_safe_title(transcript, mode_label)}.pdf"

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    except Exception:  # noqa: BLE001  已注册则忽略
        pass

    FONT = "STSong-Light"
    PAGE_W, PAGE_H = A4
    MARGIN = 1.6 * cm
    USABLE_W = PAGE_W - MARGIN * 2

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "VidTitle",
        parent=styles["Heading1"],
        fontName=FONT,
        fontSize=18,
        leading=24,
        textColor=colors.HexColor("#2c3e50"),
        spaceAfter=12,
    )
    h2_style = ParagraphStyle(
        "VidH2",
        parent=styles["Heading2"],
        fontName=FONT,
        fontSize=14,
        leading=20,
        textColor=colors.HexColor("#2c3e50"),
        spaceAfter=8,
        borderColor=colors.HexColor("#4a90d9"),
        borderWidth=1,
        borderPadding=0,
        borderBottom=True,
    )
    body_style = ParagraphStyle(
        "VidBody",
        parent=styles["BodyText"],
        fontName=FONT,
        fontSize=10,
        leading=16,
        wordWrap="CJK",
    )
    label_style = ParagraphStyle(
        "VidLabel",
        parent=body_style,
        fontName=FONT,
        textColor=colors.HexColor("#e74c3c"),
    )
    cell_style = ParagraphStyle(
        "VidCell",
        parent=body_style,
        fontSize=9,
        leading=14,
    )
    header_cell_style = ParagraphStyle(
        "VidHeaderCell",
        parent=cell_style,
        fontName=FONT,
        backColor=colors.HexColor("#f0f3f7"),
        textColor=colors.HexColor("#2c3e50"),
    )

    def _md_bold_to_html(text: str) -> str:
        """把 **...** 替换成 reportlab Paragraph 可识别的关键词高亮标记。

        与图片导出(红色 KEYWORD)、HTML 导出(strong 红色)保持一致：
        **关键词** -> <font color="#e74c3c"><b>关键词</b></font>（红色加粗）。
        STSong-Light 对 <b> 无真正粗体变体，但 <font color> 生效，关键词显红。
        """
        import re

        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return re.sub(
            r"\*\*(.+?)\*\*",
            r'<font color="#e74c3c"><b>\1</b></font>',
            text,
        )

    story = []

    # 标题
    story.append(Paragraph(_md_bold_to_html(summary.title), title_style))

    # 基本信息
    story.append(Paragraph("基本信息", h2_style))
    info_items = [
        ("来源", summary.source),
        ("作者", summary.author or "未知"),
        ("发布时间", summary.publish_time or "未知"),
        ("视频时长", summary.duration_text),
        ("内容概述", summary.content_overview or "—"),
    ]
    for label, value in info_items:
        p = Paragraph(f"<b>{label}：</b>{_md_bold_to_html(value)}", body_style)
        story.append(p)
    story.append(Spacer(1, 8))

    # 内容脉络表格
    story.append(Paragraph("内容脉络", h2_style))
    if summary.detailed:
        # 列宽：时间 12%，核心要点 20%，内容 56%，备注 12%
        col_widths = [USABLE_W * 0.12, USABLE_W * 0.20, USABLE_W * 0.56, USABLE_W * 0.12]
        data = [
            [
                Paragraph("时间", header_cell_style),
                Paragraph("核心要点", header_cell_style),
                Paragraph("内容", header_cell_style),
                Paragraph("备注", header_cell_style),
            ]
        ]
        for row in summary.detailed:
            data.append([
                Paragraph(_md_bold_to_html(row.timestamp), cell_style),
                Paragraph(_md_bold_to_html(row.point), cell_style),
                Paragraph(_md_bold_to_html(row.content), cell_style),
                Paragraph(_md_bold_to_html(row.remark or ""), cell_style),
            ])
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), FONT),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("LEADING", (0, 0), (-1, -1), 14),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f3f7")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)
    else:
        story.append(Paragraph("（暂无）", body_style))

    story.append(Spacer(1, 12))

    # 金句
    story.append(Paragraph("金句", h2_style))
    if summary.golden_quotes:
        for q in summary.golden_quotes:
            p = Paragraph(f"<b>[{q.timestamp}]</b> {_md_bold_to_html(q.text)}", body_style)
            story.append(p)
    else:
        story.append(Paragraph("（暂无）", body_style))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=MARGIN,
        leftMargin=MARGIN,
        topMargin=MARGIN,
        bottomMargin=MARGIN,
    )
    doc.build(story)
    return path


def _add_formatted_runs(paragraph, text: str) -> None:
    """把含 **加粗** 标记的文本添加到 paragraph，加粗部分设 run.bold=True。

    用于 Word 导出：python-docx 不解析 Markdown，需要手动拆分 ** 标记。
    """

    import re

    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # 奇数索引 = **...** 之间的内容 = 加粗
            run.bold = True


def export_docx(summary: Summary, output: OutputConfig, transcript: Transcript, mode_label: str = "") -> Path:
    """渲染成 Word (.docx) 文件，返回路径。"""

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()

    # 标题
    h = doc.add_heading(summary.title, level=1)

    # 基本信息 + 内容概述
    doc.add_heading("基本信息", level=2)
    info_items = [
        ("来源", summary.source),
        ("作者", summary.author or "未知"),
        ("发布时间", summary.publish_time or "未知"),
        ("视频时长", summary.duration_text),
        ("内容概述", summary.content_overview or "—"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        _add_formatted_runs(p, value)

    # 内容脉络
    doc.add_heading("内容脉络", level=2)
    if summary.detailed:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, header in enumerate(["时间", "核心要点", "内容", "备注"]):
            hdr[i].text = header
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for row in summary.detailed:
            cells = table.add_row().cells
            cells[0].text = row.timestamp
            # point 和 content 可能含 **加粗** 标记，需解析
            p1 = cells[1].paragraphs[0]
            _add_formatted_runs(p1, row.point)
            p2 = cells[2].paragraphs[0]
            _add_formatted_runs(p2, row.content)
            cells[3].text = row.remark or ""
    else:
        doc.add_paragraph("（暂无）")

    # 金句
    doc.add_heading("金句", level=2)
    if summary.golden_quotes:
        for i, q in enumerate(summary.golden_quotes, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. [{q.timestamp}] ")
            run.bold = True
            _add_formatted_runs(p, q.text)
    else:
        doc.add_paragraph("（暂无）")

    save_dir = _resolve_save_dir(output)
    save_dir.mkdir(parents=True, exist_ok=True)
    path = save_dir / f"{_safe_title(transcript, mode_label)}.docx"
    doc.save(str(path))
    return path


# ----------------------------------------------------------------------------
# 图片（PNG）导出：把摘要渲染成一张信息图风格的图，便于分享/预览
# ----------------------------------------------------------------------------

def _find_cjk_font() -> Tuple[Optional[str], Optional[str]]:
    """在常见系统中文字体里挑一个可用路径，返回 (regular_path, bold_path|None)。

    找不到返回 (None, None)。覆盖 Windows / macOS / Linux 常见中文字体。
    """
    import os

    from PIL import ImageFont

    candidates: List[Tuple[str, int, Optional[str]]] = []
    win = os.environ.get("SystemRoot", "C:/Windows")
    win_fonts = os.path.join(win, "Fonts")
    # Windows：微软雅黑（有独立粗体文件）/ 黑体 / 宋体
    candidates += [
        (os.path.join(win_fonts, "msyh.ttc"), 0, os.path.join(win_fonts, "msyhbd.ttc")),
        (os.path.join(win_fonts, "simhei.ttf"), 0, None),
        (os.path.join(win_fonts, "msyhbd.ttc"), 0, None),
        (os.path.join(win_fonts, "simsun.ttc"), 0, None),
    ]
    # macOS
    candidates += [
        ("/System/Library/Fonts/PingFang.ttc", 0, None),
        ("/System/Library/Fonts/STHeiti Light.ttc", 0, None),
        ("/System/Library/Fonts/Hiragino Sans GB.ttc", 0, None),
    ]
    # Linux
    candidates += [
        ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", 0, None),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0, None),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0, None),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", 0, None),
        ("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf", 0, None),
        ("/usr/share/fonts/truetype/arphic/uming.ttc", 0, None),
    ]
    for reg, idx, bold in candidates:
        if not os.path.exists(reg):
            continue
        try:
            ImageFont.truetype(reg, 20, index=idx if idx is not None else 0)
        except Exception:  # noqa: BLE001
            continue
        bold_path = bold if (bold and os.path.exists(bold)) else None
        return reg, bold_path
    return None, None


def export_image(summary: Summary, output: OutputConfig, transcript: Transcript, mode_label: str = "") -> Path:
    """把摘要渲染成一张 PNG 图片，布局与 HTML/Markdown 文档一致：标题/基本信息/内容脉络表格/金句列表。返回路径。

    内容脉络采用表格形式：时间 | 核心要点 | 内容 | 备注，**加粗关键词用红色高亮。
    若系统中文字体缺失，抛 FileNotFoundError，由 export() 捕获后跳过图片、不影响其他格式。
    """

    from PIL import Image, ImageDraw, ImageFont

    reg_font, bold_font = _find_cjk_font()
    if not reg_font:
        raise FileNotFoundError(
            "未找到系统中的文字体（中文），无法生成图片摘要。\n"
            "   Windows 一般自带「微软雅黑」；Linux 可装 fonts-wqy-zenhei / fonts-noto-cjk；macOS 自带 PingFang。\n"
            "   已跳过图片导出，其余格式（md/html/docx）不受影响。"
        )

    # ---- 调色板（与 HTML 样式一致）----
    BG = (255, 255, 255)
    FG = (44, 62, 80)
    MUTED = (120, 130, 140)
    ACCENT = (74, 144, 217)
    ACCENT_DARK = (44, 62, 80)
    KEYWORD = (231, 76, 60)        # 关键词红（对应 HTML 的 strong）
    TABLE_HEADER_BG = (245, 247, 250)
    TABLE_ROW_EVEN = (250, 251, 252)
    TABLE_BORDER = (221, 221, 221)
    DIVIDER = (221, 221, 221)

    W = 920
    PAD = 44
    CW = W - 2 * PAD
    TABLE_X = PAD
    TABLE_W = CW
    COLS = [90, 170, 410, 93]      # 时间 | 核心要点 | 内容 | 备注
    HEADERS = ["时间", "核心要点", "内容", "备注"]
    CELL_PAD_X = 8
    CELL_PAD_Y = 6
    BORDER = 1

    def f(sz: int, bold: bool = False) -> "ImageFont.FreeTypeFont":
        p = bold_font if (bold and bold_font) else reg_font
        return ImageFont.truetype(p, sz)

    F_TITLE = f(34, bold=True)
    F_SUB = f(18)
    F_H2 = f(26, bold=True)
    F_INFO = f(18)
    F_BODY = f(17)
    F_HEADER = f(17, bold=True)
    F_QUOTE = f(18)
    F_FOOT = f(13)

    # ---- 测量用的临时画布 ----
    meas = Image.new("RGB", (W, 20))
    dm = ImageDraw.Draw(meas)

    def text_w(t: str, font) -> float:
        return dm.textlength(t, font=font)

    def wrap_plain(t: str, font, maxw: float) -> List[str]:
        out: List[str] = []
        cur = ""
        for ch in t:
            if ch == "\n":
                out.append(cur)
                cur = ""
                continue
            if cur and text_w(cur + ch, font) > maxw:
                out.append(cur)
                cur = ch
            else:
                cur = cur + ch
        out.append(cur)
        return out

    def rich_chars(t: str):
        """把含 **加粗** 的文本拆成 [(char, is_bold), ...] 序列。"""
        seq = []
        buf = ""
        bold = False
        i = 0
        while i < len(t):
            if t[i:i + 2] == "**":
                for c in buf:
                    seq.append((c, bold))
                buf = ""
                bold = not bold
                i += 2
            else:
                buf += t[i]
                i += 1
        for c in buf:
            seq.append((c, bold))
        return seq

    def draw_rich(d, x: float, y: float, t: str, font, fill, bold_fill, maxw: float, line_h: int) -> int:
        """按宽度换行，逐字绘制（**...** 用 bold_fill 上色）。返回占用的高度。"""
        chars = rich_chars(t)
        lines: List[List[Tuple[str, bool]]] = []
        cur: List[Tuple[str, bool]] = []
        cur_w = 0.0
        for ch, b in chars:
            if ch == "\n":
                lines.append(cur)
                cur = []
                cur_w = 0.0
                continue
            w = text_w(ch, font)
            if cur and cur_w + w > maxw:
                lines.append(cur)
                cur = []
                cur_w = 0.0
            cur.append((ch, b))
            cur_w += w
        if cur:
            lines.append(cur)
        yy = y
        for line in lines:
            xx = x
            for ch, b in line:
                d.text((xx, yy), ch, font=font, fill=bold_fill if b else fill)
                xx += text_w(ch, font)
            yy += line_h
        return int(yy - y)

    # ---- 第一遍：计算布局 ----
    y = PAD

    # 标题
    title_lines = wrap_plain(summary.title or "视频摘要", F_TITLE, CW)
    title_h = len(title_lines) * (F_TITLE.size + 4)
    y += title_h + 6

    # 副标题
    y += F_SUB.size + 4
    # 分隔线
    y += 10 + 18

    # 基本信息
    y += F_H2.size + 10
    info_items = [
        ("来源", summary.source or "未知"),
        ("作者", summary.author or "未知"),
        ("发布时间", summary.publish_time or "未知"),
        ("视频时长", summary.duration_text or "未知"),
        ("内容概述", summary.content_overview or "（暂无）"),
    ]
    for label, value in info_items:
        lw = text_w(label + "：", F_INFO)
        h = draw_rich(dm, 0, 0, value, F_INFO, FG, KEYWORD, CW - lw, F_INFO.size + 8)
        y += max(h, F_INFO.size + 8) + 4
    y += 8

    # 内容脉络
    y += F_H2.size + 12

    # 表格列边界
    v_lines = [TABLE_X + i * BORDER + sum(COLS[:i]) + 2 * i * CELL_PAD_X for i in range(len(COLS) + 1)]

    # 表头
    table_header_y = y
    header_h = F_HEADER.size + 2 * CELL_PAD_Y
    y += header_h

    # 数据行
    row_specs = []
    if summary.detailed:
        for r in summary.detailed:
            contents = [r.timestamp or "", r.point or "", r.content or "", r.remark or ""]
            heights = []
            for i, content in enumerate(contents):
                h = draw_rich(dm, 0, 0, content, F_BODY, FG, KEYWORD, COLS[i], F_BODY.size + 7)
                heights.append(h)
            row_h = max(heights) + 2 * CELL_PAD_Y
            row_specs.append((r, row_h))
            y += row_h
    else:
        y += F_BODY.size + 8

    y += 8

    # 金句
    y += F_H2.size + 12
    golden_ops = []
    if summary.golden_quotes:
        for i, q in enumerate(summary.golden_quotes, 1):
            gline = f"{i}. [{q.timestamp}] {q.text}"
            h = draw_rich(dm, 0, 0, gline, F_QUOTE, FG, KEYWORD, CW, F_QUOTE.size + 7)
            golden_ops.append((gline, h))
            y += h + 8
    else:
        y += F_BODY.size + 8

    y += 6
    H = int(y + PAD)

    # ---- 第二遍：绘制 ----
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)

    # 标题
    y = PAD
    for ln in title_lines:
        d.text((PAD, y), ln, font=F_TITLE, fill=FG)
        y += F_TITLE.size + 4
    y += 6

    # 副标题
    sub = f"{summary.source or '未知'}　·　{summary.author or '未知'}　·　{summary.duration_text or '未知'}"
    d.text((PAD, y), sub, font=F_SUB, fill=MUTED)
    y += F_SUB.size + 4

    # 分隔线
    y += 10
    d.line([PAD, y, W - PAD, y], fill=DIVIDER, width=2)
    y += 18

    # 基本信息
    d.text((PAD, y), "基本信息", font=F_H2, fill=ACCENT_DARK)
    y += F_H2.size + 10
    for label, value in info_items:
        lw = text_w(label + "：", F_INFO)
        d.text((PAD, y), f"{label}：", font=F_INFO, fill=ACCENT)
        h = draw_rich(d, PAD + lw, y, value, F_INFO, FG, KEYWORD, CW - lw, F_INFO.size + 8)
        y += max(h, F_INFO.size + 8) + 4
    y += 8

    # 内容脉络
    d.text((PAD, y), "内容脉络", font=F_H2, fill=ACCENT_DARK)
    y += F_H2.size + 12

    # 表格
    if summary.detailed:
        # 表头背景
        d.rectangle(
            [v_lines[0] + BORDER, table_header_y + BORDER, v_lines[-1] - BORDER, table_header_y + header_h - BORDER],
            fill=TABLE_HEADER_BG,
        )
        # 表头文字
        for i, htext in enumerate(HEADERS):
            d.text(
                (v_lines[i] + BORDER + CELL_PAD_X, table_header_y + BORDER + CELL_PAD_Y),
                htext,
                font=F_HEADER,
                fill=FG,
            )

        # 数据行
        h_lines = [table_header_y, table_header_y + header_h]
        for _r, row_h in row_specs:
            h_lines.append(h_lines[-1] + row_h)

        row_y = table_header_y + header_h
        for idx, (r, row_h) in enumerate(row_specs):
            bg = TABLE_ROW_EVEN if idx % 2 == 0 else BG
            d.rectangle(
                [v_lines[0] + BORDER, row_y + BORDER, v_lines[-1] - BORDER, row_y + row_h - BORDER],
                fill=bg,
            )
            contents = [r.timestamp or "", r.point or "", r.content or "", r.remark or ""]
            for i, content in enumerate(contents):
                draw_rich(
                    d,
                    v_lines[i] + BORDER + CELL_PAD_X,
                    row_y + BORDER + CELL_PAD_Y,
                    content,
                    F_BODY,
                    FG,
                    KEYWORD,
                    COLS[i],
                    F_BODY.size + 7,
                )
            row_y += row_h

        # 表格边框
        for x_line in v_lines:
            d.line([(x_line, table_header_y), (x_line, h_lines[-1])], fill=TABLE_BORDER, width=BORDER)
        for y_line in h_lines:
            d.line([(TABLE_X, y_line), (TABLE_X + TABLE_W, y_line)], fill=TABLE_BORDER, width=BORDER)
        y = h_lines[-1] + 8
    else:
        d.text((PAD, y), "（暂无）", font=F_BODY, fill=MUTED)
        y += F_BODY.size + 8

    # 金句
    d.text((PAD, y), "金句", font=F_H2, fill=ACCENT_DARK)
    y += F_H2.size + 12
    if summary.golden_quotes:
        for gline, h in golden_ops:
            draw_rich(d, PAD, y, gline, F_QUOTE, FG, KEYWORD, CW, F_QUOTE.size + 7)
            y += h + 8
    else:
        d.text((PAD, y), "（暂无）", font=F_BODY, fill=MUTED)
        y += F_BODY.size + 8

    y += 6
    foot = "由 VidGrab 生成 · 带时间戳的结构化视频摘要"
    d.text((PAD, y), foot, font=F_FOOT, fill=MUTED)

    save_dir = _resolve_save_dir(output)
    save_dir.mkdir(parents=True, exist_ok=True)
    path = save_dir / f"{_safe_title(transcript, mode_label)}.png"
    img.save(str(path), "PNG")
    return path


# 格式 -> 导出函数映射
_FORMAT_EXPORTERS = {
    "markdown": export_markdown,
    "md": export_markdown,
    "html": export_html,
    "docx": export_docx,
    "word": export_docx,
    "pdf": export_pdf,
    "image": export_image,
    "png": export_image,
    "jpg": export_image,
    "jpeg": export_image,
}


def export(
    summary: Summary,
    output: Optional[OutputConfig] = None,
    transcript: Optional[Transcript] = None,
    formats: Optional[List[str]] = None,
    mode_label: str = "",
) -> List[Path]:
    """按指定格式列表导出，返回所有生成文件的路径列表。

    formats: ["markdown", "html", "docx", "image"] 的子集；None 则用 output.default_format。
    mode_label: 精简/详细/自定义，用于文件名后缀，避免互相覆盖。
    单个格式导出失败（如缺少中文字体导致图片失败）不影响其他格式。
    """

    output = output or OutputConfig()
    transcript = transcript or Transcript(platform=summary.platform, video_id="output")  # type: ignore

    if formats:
        fmt_list = formats
    else:
        fmt_list = [(output.default_format or "markdown").lower()]

    paths: List[Path] = []
    for fmt in fmt_list:
        fmt = fmt.lower().strip()
        exporter_fn = _FORMAT_EXPORTERS.get(fmt)
        if not exporter_fn:
            print(f"⚠️ 不支持的格式：{fmt}（支持 markdown/html/docx/pdf/image），跳过")
            continue
        try:
            paths.append(exporter_fn(summary, output, transcript, mode_label))
        except Exception as e:  # noqa: BLE001
            print(f"⚠️ 导出 {fmt} 失败：{e}（其余格式不受影响）")
    return paths
